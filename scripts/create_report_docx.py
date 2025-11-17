"""Generate a structured Word report with placeholders for figures and tables."""
from pathlib import Path

from docx import Document


def add_section_heading(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def add_paragraph(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_bullets(doc: Document, items) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_placeholder(doc: Document, label: str, description: str) -> None:
    doc.add_paragraph(f"{label}: {description}", style="Intense Quote")


def main() -> None:
    doc = Document()
    doc.add_heading("NBA Prediction Project – Draft Report", level=1)

    # Introduction
    add_section_heading(doc, "Introduction")
    add_paragraph(
        doc,
        (
            "This study leverages the 2004–2005 Basketball Reference snapshot housed entirely in the local "
            "databasebasketball folder. We pursue two objectives: (1) identify outstanding players via robust "
            "outlier detection, and (2) predict hypothetical game outcomes between any two teams using season-level "
            "statistics. The workflow emphasises reproducibility by relying solely on local data, Python modules in src/, "
            "and Jupyter notebooks for analysis."
        ),
    )

    # Methods
    add_section_heading(doc, "Methods and Techniques")
    method_points = [
        "Data ingestion: src/data_loading.py standardises column names/dtypes for players, career totals, playoff stats, "
        "all-star data, and team-season totals.",
        "Player-level features: src/feature_engineering.build_player_feature_table joins players.txt with regular-season, "
        "playoff, and all-star stats via ilkid, filters careers with fewer than 82 games, and engineers per-game/per-36 "
        "rates, efficiency percentages, impact_score, usage_proxy, and postseason indicators.",
        "Outlier detection: Notebook 01 applies MAD-based robust z-scores and IsolationForest via src.models_outliers, "
        "exporting tables/top_outliers_mad.csv, tables/top_outliers_iforest.csv, tables/outlier_overlap.csv, and figures "
        "such as figures/player_metric_distributions.png, figures/player_metric_correlation.png, and "
        "figures/usage_vs_impact_outliers.png.",
        "Team features: src/feature_engineering.build_team_season_features creates per-game offense/defense, margin, "
        "pace-adjusted ratings, turnover/assist/rebound rates, and readable team names; build_pairwise_matchups forms "
        "ordered team pairs with difference features and win%-based labels.",
        "Game outcome modelling: Notebook 02 performs season-aware splits, trains logistic regression, gradient boosting, "
        "and random forest models (src.models_game_outcome), and evaluates them with src.evaluation to produce "
        "tables/model_comparison.csv plus figures/confusion_*.png, figures/roc_models.png, and "
        "figures/feature_importance_random_forest.png.",
        "Results consolidation: Notebook 03 loads the saved artifacts and assembles report-ready tables, figures, and "
        "narrative bullets."
    ]
    add_bullets(doc, method_points)

    # Results
    add_section_heading(doc, "Results and Discussion")
    result_points = [
        "Outstanding players: MAD surfaces statistically peculiar stat lines (see Table 1 placeholder), whereas "
        "IsolationForest highlights multi-dimensional legends (Table 2 placeholder). The empty overlap table underscores "
        "that the two methods encode different definitions of “outstanding.” Refer to Figure 1 for the usage vs. impact "
        "scatter with MAD outliers annotated.",
        "Game outcome models: Difference features derived from season stats yielded near-perfect separation. Logistic "
        "regression obtained accuracy 0.999/ROC-AUC 1.0, while gradient boosting and random forest were effectively "
        "perfect (Table 3 placeholder). Confusion matrices (Figure 2), ROC curves (Figure 3), and feature-importance "
        "plots (Figure 4) illustrate the models’ behaviour and the dominance of diff_win_pct and diff_margin."
    ]
    add_bullets(doc, result_points)

    # Placeholders
    add_section_heading(doc, "Figures and Tables to Insert")
    add_placeholder(doc, "Table 1", "Top MAD outliers (tables/top_outliers_mad.csv).")
    add_placeholder(doc, "Table 2", "Top IsolationForest outliers (tables/top_outliers_iforest.csv).")
    add_placeholder(doc, "Table 3", "Model comparison metrics (tables/model_comparison.csv).")
    add_placeholder(doc, "Figure 1", "Usage vs. impact scatter with MAD flags (figures/usage_vs_impact_outliers.png).")
    add_placeholder(doc, "Figure 2", "Confusion matrices for logistic regression and gradient boosting "
                                     "(figures/confusion_logistic_regression.png, figures/confusion_gradient_boosting.png).")
    add_placeholder(doc, "Figure 3", "ROC curves for all models (figures/roc_models.png).")
    add_placeholder(doc, "Figure 4", "Random forest feature importances (figures/feature_importance_random_forest.png).")

    # Conclusion
    add_section_heading(doc, "Conclusion")
    add_paragraph(
        doc,
        (
            "The project delivers an end-to-end ML pipeline: modular loaders, dual outlier analyses, and interpretable "
            "matchup models with reproducible figures/tables. MAD and IsolationForest provide complementary views on "
            "outstanding players, while matchup models confirm that season aggregates encode enough signal to rank "
            "teams reliably. Future extensions could incorporate era-adjusted features or actual game outcomes, and "
            "the report must cite basketballreference.com per the dataset license."
        ),
    )

    output = Path("report_draft.docx")
    doc.save(output)
    print(f"Saved {output.resolve()}")


if __name__ == "__main__":
    main()
